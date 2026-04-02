"""
Generador del Manual de Administrador - Sistema PQRS Conjunto Calle 100
Ejecutar: py -3 gen_manual_admin.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ─── Configuración de página ───────────────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin   = Inches(1.2)
section.right_margin  = Inches(1.2)
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)

# ─── Estilos base (RGBColor para fuentes, hex para fondos de celdas) ───────
VERDE_OSCURO = RGBColor(0x1B, 0x5E, 0x20)
VERDE_MEDIO  = RGBColor(0x2E, 0x7D, 0x32)
AZUL_TEXTO   = RGBColor(0x0D, 0x47, 0xA1)
NARANJA      = RGBColor(0xFF, 0x6F, 0x00)
ROJO_TEXTO   = RGBColor(0xB7, 0x1C, 0x1C)
GRIS_OSCURO  = RGBColor(0x42, 0x42, 0x42)
BLANCO       = RGBColor(0xFF, 0xFF, 0xFF)

# Hex strings para fondos de celda
H_VERDE_OSCURO = "1B5E20"
H_VERDE_MEDIO  = "2E7D32"
H_VERDE_CLARO  = "E8F5E9"
H_AZUL_CLARO   = "E3F2FD"
H_NARANJA_FONDO= "FFF8E1"
H_ROJO_FONDO   = "FFEBEE"
H_GRIS_FONDO   = "F5F5F5"
H_BLANCO       = "FFFFFF"

FONT = "Calibri"

# ─── Helpers ───────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """hex_color: string hex p.ej. 'E8F5E9' (sin #)"""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color.lstrip("#"))
    tcPr.append(shd)

def set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, color in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if color:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"),  "6")
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tcBorders.append(el)
    tcPr.append(tcBorders)

def para_style(para, font_name=FONT, size=11, bold=False, color=None, align=None, space_before=0, space_after=6):
    run = para.runs[0] if para.runs else para.add_run("")
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    if align:
        para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after  = Pt(space_after)
    return para

def add_page_break():
    doc.add_page_break()

def heading1(text, numbering=""):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "1B5E20")
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(f"{numbering}  {text}" if numbering else text)
    run.font.name  = FONT
    run.font.size  = Pt(16)
    run.font.bold  = True
    run.font.color.rgb = VERDE_OSCURO
    return p

def heading2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.font.name  = FONT
    run.font.size  = Pt(13)
    run.font.bold  = True
    run.font.color.rgb = VERDE_MEDIO
    return p

def heading3(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = FONT
    run.font.size  = Pt(11)
    run.font.bold  = True
    run.font.color.rgb = GRIS_OSCURO
    return p

def body(text, size=10.5, color=None, bold=False, space_after=5):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.font.name  = FONT
    run.font.size  = Pt(size)
    run.font.bold  = bold
    if color:
        run.font.color.rgb = color
    return p

def bullet(text, level=0, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent   = Inches(0.3 + level * 0.25)
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.space_before  = Pt(0)
    if bold_prefix:
        r = p.add_run(bold_prefix + ": ")
        r.font.name  = FONT
        r.font.size  = Pt(10)
        r.font.bold  = True
        r.font.color.rgb = VERDE_MEDIO
        r2 = p.add_run(text)
        r2.font.name  = FONT
        r2.font.size  = Pt(10)
    else:
        r = p.add_run(text)
        r.font.name  = FONT
        r.font.size  = Pt(10)
    return p

def note_box(text, title="Nota"):
    """Caja de nota verde."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, H_VERDE_CLARO)
    set_cell_border(cell, top="2E7D32", bottom="2E7D32", left="2E7D32", right="2E7D32")
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"{title}: ")
    r1.font.name = FONT; r1.font.size = Pt(9.5); r1.font.bold = True; r1.font.color.rgb = VERDE_OSCURO
    r2 = p1.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(9.5); r2.font.color.rgb = VERDE_OSCURO
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def warning_box(text, title="Importante"):
    """Caja de advertencia naranja."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, H_NARANJA_FONDO)
    set_cell_border(cell, top="FF6F00", bottom="FF6F00", left="FF6F00", right="FF6F00")
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"⚠ {title}: ")
    r1.font.name = FONT; r1.font.size = Pt(9.5); r1.font.bold = True; r1.font.color.rgb = NARANJA
    r2 = p1.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(9.5); r2.font.color.rgb = GRIS_OSCURO
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def info_box(text, title="Info"):
    """Caja de información azul."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, H_AZUL_CLARO)
    set_cell_border(cell, top="0D47A1", bottom="0D47A1", left="0D47A1", right="0D47A1")
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(f"ℹ {title}: ")
    r1.font.name = FONT; r1.font.size = Pt(9.5); r1.font.bold = True; r1.font.color.rgb = AZUL_TEXTO
    r2 = p1.add_run(text)
    r2.font.name = FONT; r2.font.size = Pt(9.5); r2.font.color.rgb = GRIS_OSCURO
    p1.paragraph_format.space_before = Pt(2)
    p1.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return tbl

def step_table(steps):
    """Tabla de pasos numerados."""
    tbl = doc.add_table(rows=len(steps), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Inches(0.45), Inches(4.7)]
    for i, (num, text) in enumerate(steps):
        c0, c1 = tbl.cell(i, 0), tbl.cell(i, 1)
        set_cell_bg(c0, H_VERDE_MEDIO)
        set_cell_bg(c1, H_GRIS_FONDO if i % 2 == 0 else H_BLANCO)
        c0.width = col_widths[0]
        c1.width = col_widths[1]
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(str(num))
        r0.font.name = FONT; r0.font.size = Pt(11); r0.font.bold = True; r0.font.color.rgb = BLANCO
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.paragraph_format.space_before = Pt(2); p0.paragraph_format.space_after = Pt(2)
        p1 = c1.paragraphs[0]
        if isinstance(text, tuple):
            rb = p1.add_run(text[0] + ": ")
            rb.font.name = FONT; rb.font.size = Pt(10); rb.font.bold = True; rb.font.color.rgb = VERDE_OSCURO
            rn = p1.add_run(text[1])
            rn.font.name = FONT; rn.font.size = Pt(10)
        else:
            r1 = p1.add_run(text)
            r1.font.name = FONT; r1.font.size = Pt(10)
        p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

def field_table(rows_data, header=None):
    """Tabla de campos/descripción."""
    n = len(rows_data)
    has_header = header is not None
    tbl = doc.add_table(rows=n + (1 if has_header else 0), cols=2)
    tbl.style = "Table Grid"
    col_widths = [Inches(1.8), Inches(3.35)]
    offset = 0
    if has_header:
        c0, c1 = tbl.cell(0, 0), tbl.cell(0, 1)
        set_cell_bg(c0, H_VERDE_OSCURO); set_cell_bg(c1, H_VERDE_OSCURO)
        for c, t in [(c0, header[0]), (c1, header[1])]:
            p = c.paragraphs[0]
            r = p.add_run(t)
            r.font.name = FONT; r.font.size = Pt(10); r.font.bold = True; r.font.color.rgb = BLANCO
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        offset = 1
    for i, (field, desc) in enumerate(rows_data):
        c0, c1 = tbl.cell(i + offset, 0), tbl.cell(i + offset, 1)
        c0.width = col_widths[0]; c1.width = col_widths[1]
        set_cell_bg(c0, H_VERDE_CLARO if i % 2 == 0 else H_BLANCO)
        set_cell_bg(c1, H_GRIS_FONDO if i % 2 == 0 else H_BLANCO)
        p0 = c0.paragraphs[0]
        r0 = p0.add_run(field)
        r0.font.name = FONT; r0.font.size = Pt(9.5); r0.font.bold = True; r0.font.color.rgb = VERDE_OSCURO
        p0.paragraph_format.space_before = Pt(2); p0.paragraph_format.space_after = Pt(2)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(desc)
        r1.font.name = FONT; r1.font.size = Pt(9.5)
        p1.paragraph_format.space_before = Pt(2); p1.paragraph_format.space_after = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return tbl

# ═══════════════════════════════════════════════════════════════════════════
# PORTADA
# ═══════════════════════════════════════════════════════════════════════════

# Franja verde superior simulada con tabla
tbl_cover = doc.add_table(rows=1, cols=1)
tbl_cover.style = "Table Grid"
c = tbl_cover.cell(0, 0)
set_cell_bg(c, H_VERDE_OSCURO)
p = c.paragraphs[0]
r = p.add_run("SISTEMA PQRS — CONJUNTO CALLE 100")
r.font.name = FONT; r.font.size = Pt(20); r.font.bold = True; r.font.color.rgb = BLANCO
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(14)

doc.add_paragraph()

p_title = doc.add_paragraph()
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_title.add_run("MANUAL DE ADMINISTRADOR")
r.font.name = FONT; r.font.size = Pt(28); r.font.bold = True; r.font.color.rgb = VERDE_OSCURO
p_title.paragraph_format.space_before = Pt(60)
p_title.paragraph_format.space_after  = Pt(8)

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_sub.add_run("Guía Completa del Sistema de Gestión de Peticiones,\nQuejas, Reclamos y Sugerencias")
r.font.name = FONT; r.font.size = Pt(13); r.font.color.rgb = GRIS_OSCURO
p_sub.paragraph_format.space_after = Pt(60)

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

p_meta = doc.add_paragraph()
p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_meta.add_run("Versión 1.0  ·  Año 2026\nUso interno — Administración, Asistentes y Consejo de Administración")
r.font.name = FONT; r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0x75, 0x75, 0x75)
p_meta.paragraph_format.space_after = Pt(4)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# TABLA DE CONTENIDO
# ═══════════════════════════════════════════════════════════════════════════

heading1("TABLA DE CONTENIDO")

toc_items = [
    ("1", "Introducción y Roles del Sistema", "4"),
    ("2", "Acceso a la Plataforma", "5"),
    ("3", "Panel de Control (Dashboard)", "6"),
    ("4", "Módulo PQRS — Lista y Búsqueda", "9"),
    ("5", "Crear PQRS Manualmente", "11"),
    ("6", "Detalle de una PQRS — Gestión Completa", "12"),
    ("7", "Sistema de Fases de Gestión", "16"),
    ("8", "Cierre y Evidencias", "22"),
    ("9", "Módulo de Reportes", "24"),
    ("10", "Gestión de Usuarios", "26"),
    ("11", "Preguntas Frecuentes", "28"),
]

tbl_toc = doc.add_table(rows=len(toc_items), cols=2)
tbl_toc.style = "Table Grid"
for i, (num, title, pg) in enumerate(toc_items):
    c0, c1 = tbl_toc.cell(i, 0), tbl_toc.cell(i, 1)
    set_cell_bg(c0, H_VERDE_CLARO if i % 2 == 0 else H_BLANCO)
    set_cell_bg(c1, H_GRIS_FONDO if i % 2 == 0 else H_BLANCO)
    p0 = c0.paragraphs[0]
    rn = p0.add_run(f"{num}.")
    rn.font.name = FONT; rn.font.size = Pt(10); rn.font.bold = True; rn.font.color.rgb = VERDE_OSCURO
    rt = p0.add_run(f"  {title}")
    rt.font.name = FONT; rt.font.size = Pt(10)
    p0.paragraph_format.space_before = Pt(3); p0.paragraph_format.space_after = Pt(3)
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(pg)
    r1.font.name = FONT; r1.font.size = Pt(10); r1.font.color.rgb = GRIS_OSCURO
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(3); p1.paragraph_format.space_after = Pt(3)

doc.add_paragraph()
add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 1: INTRODUCCIÓN Y ROLES
# ═══════════════════════════════════════════════════════════════════════════

heading1("1. Introducción y Roles del Sistema", "")
body(
    "El Sistema PQRS del Conjunto Calle 100 es una plataforma web centralizada que permite a los residentes "
    "presentar Peticiones, Quejas, Reclamos y Sugerencias, y a la administración gestionarlas de forma "
    "estructurada con trazabilidad completa. Este manual está dirigido a todos los perfiles con acceso "
    "administrativo: Administrador, Asistente y Consejo de Administración."
)

heading2("1.1 Roles del Sistema")

field_table(
    [
        ("ADMIN", "Acceso total. Puede crear PQRS a nombre de residentes, gestionar todas las fases, "
                  "asignar estados, generar reportes, administrar usuarios y ver todas las secciones."),
        ("ASISTENTE", "Acceso a gestión de PQRS: puede registrar primer contacto, acción tomada y "
                      "avanzar fases. No puede crear usuarios ni acceder a reportes globales."),
        ("CONSEJO", "Acceso de solo lectura. Puede ver el listado de PQRS activas, el detalle de "
                    "cada PQRS y las fases de gestión. No puede editar ni gestionar."),
        ("RESIDENTE", "Acceso restringido. Solo ve sus propias PQRS y puede crear nuevas. "
                      "No ve fases internas ni información de gestión detallada."),
    ],
    header=("Rol", "Permisos y Acceso")
)

note_box(
    "El Administrador es el único perfil que puede registrar y dar de baja a usuarios, "
    "asignar roles, y acceder al módulo de reportes completo.",
    title="Nota"
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 2: ACCESO A LA PLATAFORMA
# ═══════════════════════════════════════════════════════════════════════════

heading1("2. Acceso a la Plataforma")
body("El acceso se realiza mediante usuario y contraseña. No existe autenticación por enlace mágico ni registro abierto para perfiles administrativos.")

heading2("2.1 Iniciar Sesión")
step_table([
    (1, "Abra el navegador web e ingrese la URL del sistema."),
    (2, "Ingrese su correo electrónico institucional en el campo 'Correo electrónico'."),
    (3, "Ingrese su contraseña en el campo correspondiente."),
    (4, "Haga clic en el botón 'Ingresar'. El sistema lo redirigirá automáticamente al Panel de Control."),
])

warning_box(
    "Si olvida su contraseña, contacte al Administrador del sistema para que la restablezca. "
    "No existe función de recuperación automática para perfiles administrativos.",
    title="Contraseña olvidada"
)

heading2("2.2 Cerrar Sesión")
body("Para cerrar sesión de forma segura:")
step_table([
    (1, "Haga clic en su nombre o foto de perfil ubicado en la esquina superior derecha de la pantalla."),
    (2, ("Seleccione", "la opción 'Cerrar sesión' del menú desplegable.")),
    (3, "El sistema cerrará su sesión y lo redirigirá a la pantalla de inicio de sesión."),
])
warning_box("Siempre cierre sesión al terminar de usar el sistema, especialmente en equipos compartidos.")

heading2("2.3 Cambiar Contraseña")
step_table([
    (1, "Haga clic en su nombre en la esquina superior derecha."),
    (2, "Seleccione 'Cambiar contraseña'."),
    (3, "Ingrese su contraseña actual y luego la nueva contraseña (mínimo 8 caracteres)."),
    (4, "Confirme la nueva contraseña y haga clic en 'Guardar'."),
])

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 3: PANEL DE CONTROL (DASHBOARD)
# ═══════════════════════════════════════════════════════════════════════════

heading1("3. Panel de Control (Dashboard)")
body(
    "El Dashboard es la pantalla principal del sistema. Ofrece una vista consolidada del estado actual "
    "de todas las PQRS del conjunto, con indicadores clave, tablas de seguimiento y accesos directos "
    "a las secciones más importantes."
)

heading2("3.1 Tarjetas de Resumen")
body("En la parte superior del Dashboard se muestran tarjetas con las métricas más importantes:")

field_table(
    [
        ("Total PQRS",       "Número total de PQRS registradas en el sistema (todos los estados)."),
        ("En Espera",        "PQRS recibidas que aún no tienen gestión iniciada."),
        ("En Proceso",       "PQRS con gestión activa: tienen primer contacto registrado o fases iniciadas."),
        ("Terminadas",       "PQRS cerradas con evidencia de cierre."),
        ("Tiempo Promedio",  "Promedio de días entre la fecha de recibo y la fecha de cierre de PQRS terminadas."),
    ],
    header=("Indicador", "Descripción")
)

heading2("3.2 Tabla de PQRS Activas")
body(
    "La tabla principal del Dashboard muestra todas las PQRS en estado EN_ESPERA y EN_PROGRESO. "
    "Esta vista permite identificar rápidamente los casos pendientes y su antigüedad."
)
body("Columnas de la tabla:")
bullet("N°: Número único de radicación de la PQRS.")
bullet("Residente: Nombre del residente que presentó la PQRS.")
bullet("Bloque / Apto: Ubicación del residente en el conjunto.")
bullet("Asunto: Categoría de la PQRS (área común, convivencia, humedad, etc.).")
bullet("Fecha Recibido: Fecha en que fue registrada la PQRS.")
bullet("Estado: EN ESPERA (amarillo) o EN PROCESO (azul).")
bullet("Días: Número de días transcurridos desde que fue recibida.")

note_box(
    "Haga clic sobre cualquier fila de la tabla para abrir el detalle completo de esa PQRS.",
    title="Acceso rápido"
)

heading2("3.3 Tablas por Trimestre")
body(
    "Debajo de la tabla principal aparecen tablas agrupadas por trimestre del año. Cada tabla "
    "muestra el resumen de PQRS por asunto en ese período, con totales por columna."
)
body("Trimestres del año:")
field_table(
    [
        ("I Trimestre",   "Enero, Febrero, Marzo"),
        ("II Trimestre",  "Abril, Mayo, Junio"),
        ("III Trimestre", "Julio, Agosto, Septiembre"),
        ("IV Trimestre",  "Octubre, Noviembre, Diciembre"),
    ]
)

heading2("3.4 Tabla de PQRS por Bloque")
body(
    "Esta sección muestra una tabla con el total de PQRS recibidas agrupadas por bloque (torre) del conjunto. "
    "Permite identificar qué bloques tienen más solicitudes activas."
)

heading2("3.5 Accesos Directos del Dashboard")
body("Desde el Dashboard puede navegar directamente a:")
bullet("Ver todas las PQRS activas → sección PQRS")
bullet("Ver PQRS en espera → filtra automáticamente por estado EN_ESPERA")
bullet("Ver PQRS en proceso → filtra automáticamente por estado EN_PROGRESO")
bullet("Ver PQRS terminadas → filtra automáticamente por estado TERMINADO")

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 4: MÓDULO PQRS — LISTA Y BÚSQUEDA
# ═══════════════════════════════════════════════════════════════════════════

heading1("4. Módulo PQRS — Lista y Búsqueda")
body(
    "El módulo PQRS contiene el listado completo de todas las solicitudes registradas. "
    "Dispone de filtros avanzados para encontrar PQRS específicas."
)

heading2("4.1 Vista Predeterminada")
body(
    "Al ingresar al módulo PQRS sin ningún filtro activo, el sistema muestra por defecto las PQRS "
    "activas (EN ESPERA + EN PROCESO), ordenadas de más reciente a más antigua. "
    "Esto permite enfocarse en los casos pendientes de gestión."
)

heading2("4.2 Filtros Disponibles")
field_table(
    [
        ("Estado",    "Filtra por estado: Activas (predeterminado), Todas, En Espera, En Proceso, Terminadas."),
        ("Mes",       "Filtra PQRS por el mes en que fueron recibidas."),
        ("Año",       "Filtra PQRS por el año de recepción."),
        ("Asunto",    "Filtra por categoría: Área Común, Contabilidad, Convivencia, Humedad (5 subcategorías)."),
        ("N° PQRS",   "Búsqueda exacta por número de radicación (p.ej.: 47)."),
        ("Bloque",    "Filtra todas las PQRS de un bloque específico (1 al 12)."),
        ("Apto",      "Filtra PQRS de un apartamento específico."),
    ],
    header=("Filtro", "Descripción")
)

info_box(
    "Puede combinar múltiples filtros simultáneamente. Por ejemplo: Asunto='HUMEDAD/CUBIERTA' + Año='2026' "
    "para ver todas las PQRS de humedad en cubierta del año actual.",
    title="Filtros combinados"
)

body("Para limpiar todos los filtros, haga clic en el botón 'Limpiar' que aparece cuando hay filtros activos.")

heading2("4.3 Tarjetas de PQRS en la Lista")
body("Cada PQRS en la lista muestra:")
bullet("Número de radicación (p.ej.: #47)")
bullet("Estado actual con color indicador")
bullet("Asunto (o primeras palabras de la descripción si no tiene asunto asignado)")
bullet("Nombre del residente, bloque y apartamento")
bullet("Fecha de recepción")

heading2("4.4 Conteo de Resultados")
body(
    "Sobre la lista se muestra el número total de solicitudes que cumplen los filtros activos. "
    "Ejemplo: '23 solicitudes'."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 5: CREAR PQRS MANUALMENTE
# ═══════════════════════════════════════════════════════════════════════════

heading1("5. Crear PQRS Manualmente")
body(
    "El Administrador puede crear PQRS a nombre de un residente cuando la solicitud llega por "
    "medios fuera de la plataforma (correo electrónico, comunicación verbal, WhatsApp, etc.)."
)

warning_box(
    "Solo el perfil ADMIN puede crear PQRS manualmente. Los perfiles ASISTENTE y CONSEJO no tienen "
    "esta opción.",
    title="Restricción"
)

heading2("5.1 Pasos para Crear una PQRS")
step_table([
    (1, "En el módulo PQRS, haga clic en el botón verde 'Nueva' en la esquina superior derecha."),
    (2, "Complete el campo 'Nombre del residente' con el nombre completo."),
    (3, "Seleccione el 'Bloque' (1–12) y el número de 'Apartamento'."),
    (4, "Seleccione el 'Asunto' correspondiente del listado desplegable (opcional pero recomendado)."),
    (5, "Escriba la descripción detallada de la solicitud en el campo 'Descripción' (máximo 300 palabras)."),
    (6, "Haga clic en 'Crear PQRS'. El sistema asignará automáticamente el número de radicación."),
])

heading2("5.2 Campos del Formulario")
field_table(
    [
        ("Nombre del residente",  "Nombre completo de quien presenta la solicitud. Obligatorio."),
        ("Bloque",               "Número de bloque/torre (1 al 12). Obligatorio."),
        ("Apartamento",          "Número de apartamento. Obligatorio."),
        ("Asunto",               "Categoría de la solicitud. Opcional, puede asignarse después."),
        ("Descripción",          "Texto completo de la solicitud. Obligatorio. Máximo 300 palabras. "
                                  "Se muestra un contador de palabras en tiempo real."),
    ],
    header=("Campo", "Descripción")
)

note_box(
    "Después de crear la PQRS, el sistema la asigna automáticamente al estado EN ESPERA "
    "y queda disponible en la lista general para gestión.",
    title="Estado inicial"
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 6: DETALLE DE UNA PQRS — GESTIÓN COMPLETA
# ═══════════════════════════════════════════════════════════════════════════

heading1("6. Detalle de una PQRS — Gestión Completa")
body(
    "Al hacer clic sobre una PQRS en la lista, se abre la vista de detalle. Esta es la pantalla "
    "central de gestión donde se registra toda la actividad sobre la solicitud."
)

heading2("6.1 Información de la PQRS")
body("En la parte superior del detalle se muestra la información de la solicitud:")
bullet("Número de radicación y estado actual")
bullet("Tipo de PQRS (Petición, Queja, Reclamo, Sugerencia) — si fue asignado")
bullet("Asunto y sub-asunto — si fueron asignados")
bullet("Nombre del residente, bloque y apartamento")
bullet("Fecha de recepción y mes")
bullet("Descripción completa de la solicitud")
bullet("Usuario que creó la PQRS (residente o administración)")

heading2("6.2 Panel de Gestión (Admin / Asistente)")
body(
    "Debajo de la información de la PQRS, el Administrador y el Asistente ven el panel de gestión "
    "con los siguientes campos editables:"
)

field_table(
    [
        ("Tipo de PQRS",             "Clasifica la solicitud: Petición, Queja, Reclamo o Sugerencia."),
        ("Asunto",                   "Categoría de la solicitud (Área Común, Convivencia, Humedad, etc.)."),
        ("Sub-asunto",               "Detalle adicional de la categoría (campo libre)."),
        ("Primer Contacto",          "Fecha en que se comunicó al residente por primera vez. "
                                      "Al guardar, se calcula automáticamente el tiempo de respuesta inicial."),
        ("Nota Primer Contacto",     "Detalle de la comunicación inicial con el residente."),
        ("Acción Tomada",            "Descripción de las acciones realizadas para resolver la solicitud. "
                                      "Solo visible para perfiles administrativos."),
        ("Estado",                   "Cambia el estado de la PQRS: EN ESPERA, EN PROCESO o TERMINADO."),
    ],
    header=("Campo", "Descripción")
)

note_box(
    "Todos los cambios guardados quedan registrados en el Historial de la PQRS con fecha, "
    "estado anterior, estado nuevo y notas.",
    title="Trazabilidad"
)

heading2("6.3 Cómo Registrar el Primer Contacto")
step_table([
    (1, "Abra el detalle de la PQRS."),
    (2, "En el campo 'Primer Contacto', seleccione la fecha en que se comunicó al residente."),
    (3, "Escriba una nota describiendo la comunicación (p.ej.: 'Se llamó al residente el 15/03/2026, se explicó el proceso')."),
    (4, "Cambie el estado a 'EN PROCESO' si la gestión ya fue iniciada."),
    (5, "Haga clic en 'Guardar'. El sistema calculará el número de días entre la fecha de recibo y el primer contacto."),
])

heading2("6.4 Registro de Acción Tomada")
body(
    "El campo 'Acción Tomada' permite documentar todas las gestiones realizadas para resolver la PQRS. "
    "Es un campo de texto libre donde se puede escribir el historial de acciones. "
    "Este campo solo es visible para Administrador, Asistente y Consejo; no lo ve el residente."
)

heading2("6.5 Historial de Cambios")
body(
    "Al final del detalle de cada PQRS se muestra el historial de todos los cambios de estado, "
    "con fecha, estado anterior, estado nuevo y nota del cambio."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 7: SISTEMA DE FASES DE GESTIÓN
# ═══════════════════════════════════════════════════════════════════════════

heading1("7. Sistema de Fases de Gestión")
body(
    "El sistema de fases permite documentar el progreso detallado de una PQRS a través de "
    "etapas de trabajo definidas. Las fases son visibles solo para Administrador, Asistente y Consejo. "
    "El residente no puede ver las fases internas; solo ve el estado general."
)

warning_box(
    "Las fases de gestión solo están disponibles para PQRS que se encuentran en estado EN PROCESO. "
    "Una PQRS EN ESPERA no tiene fases activas.",
    title="Requisito"
)

heading2("7.1 Las 5 Fases del Proceso")

field_table(
    [
        ("Fase I — Inspección de Campo",
         "Primera fase obligatoria. El equipo técnico realiza una inspección física para "
         "evaluar el problema. Se registra la fecha de inicio y una nota descriptiva."),
        ("Fase II — Adquisición de Insumos",
         "Aplica cuando el problema requiere compra de materiales. "
         "Se activa solo si el Tipo de Fase es 'INSUMOS'. "
         "Se registra fecha de inicio y nota."),
        ("Fase III — Firma Contrato Proveedor",
         "Aplica cuando el problema requiere contratar un proveedor externo. "
         "Se activa solo si el Tipo de Fase es 'PROVEEDOR'. "
         "Solo una de las fases II o III se usa por PQRS."),
        ("Fase IV — Ejecución",
         "Fase de ejecución de la solución (trabajo en campo, reparación, instalación, etc.). "
         "Se registra fecha de inicio y nota descriptiva de lo que se está haciendo. "
         "La nota de esta fase es la que puede ver el residente como 'En ejecución'."),
        ("Fase V — Terminado",
         "Fase final. Indica que el trabajo fue completado. "
         "Al completar la Fase V, el sistema sugiere cerrar la PQRS."),
    ],
    header=("Fase", "Descripción")
)

heading2("7.2 Tipo de Fase (INSUMOS vs PROVEEDOR)")
body(
    "Antes o durante la gestión, se debe seleccionar el 'Tipo de Fase' que determina el camino de resolución:"
)

tbl_tipo = doc.add_table(rows=3, cols=2)
tbl_tipo.style = "Table Grid"
headers = [("Tipo de Fase", "Ruta de Fases"), ("INSUMOS", "Fase I → Fase II → Fase IV → Fase V"), ("PROVEEDOR", "Fase I → Fase III → Fase IV → Fase V")]
for i, (t, r) in enumerate(headers):
    c0, c1 = tbl_tipo.cell(i, 0), tbl_tipo.cell(i, 1)
    if i == 0:
        set_cell_bg(c0, H_VERDE_OSCURO); set_cell_bg(c1, H_VERDE_OSCURO)
        for c, tx in [(c0, t), (c1, r)]:
            p = c.paragraphs[0]
            rr = p.add_run(tx)
            rr.font.name = FONT; rr.font.size = Pt(10); rr.font.bold = True; rr.font.color.rgb = BLANCO
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
    else:
        set_cell_bg(c0, H_VERDE_CLARO if i % 2 else H_BLANCO)
        set_cell_bg(c1, H_GRIS_FONDO if i % 2 else H_BLANCO)
        for c, tx, bld in [(c0, t, True), (c1, r, False)]:
            p = c.paragraphs[0]
            rr = p.add_run(tx)
            rr.font.name = FONT; rr.font.size = Pt(10); rr.font.bold = bld
            if bld: rr.font.color.rgb = VERDE_OSCURO
            p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

info_box(
    "Si el problema es de bajo costo y la administración lo resuelve directamente (sin insumos ni proveedor), "
    "puede omitir la selección de tipo y avanzar directamente de Fase I a Fase IV.",
    title="Caso especial"
)

heading2("7.3 Indicador de Semáforo")
body(
    "Cada fase activa muestra un indicador de tiempo basado en los días transcurridos desde el inicio de esa fase:"
)
field_table(
    [
        ("Verde",    "0 a 7 días — La fase está dentro del tiempo esperado."),
        ("Amarillo", "8 a 14 días — La fase está próxima a vencer."),
        ("Rojo",     "15 días o más — La fase ha excedido el tiempo recomendado."),
    ],
    header=("Color", "Significado")
)

heading2("7.4 Avanzar a la Siguiente Fase")
step_table([
    (1, "Abra el detalle de una PQRS en estado EN PROCESO."),
    (2, "En la sección 'Fases de Gestión', ubique la fase actual (indicada con el semáforo de color)."),
    (3, "Escriba la nota de la fase en el campo de texto correspondiente. La nota es obligatoria para avanzar."),
    (4, "Haga clic en el botón 'Avanzar a Fase X' donde X es la siguiente fase."),
    (5, "El sistema registrará la fecha de inicio de la nueva fase y mostrará el progreso actualizado."),
])

warning_box(
    "No es posible avanzar a la siguiente fase sin escribir la nota de la fase actual. "
    "La nota documenta qué se hizo en esa etapa.",
    title="Nota obligatoria"
)

heading2("7.5 Visualización para el Consejo")
body(
    "El perfil CONSEJO puede ver las fases de gestión de las PQRS en estado EN PROCESO y TERMINADO, "
    "pero en modo de solo lectura. No puede editar notas ni avanzar fases."
)
body("El Consejo verá:")
bullet("Las fases completadas con fecha de inicio, días transcurridos y nota registrada.")
bullet("El semáforo de tiempo de la fase activa.")
bullet("El tipo de fase seleccionado (INSUMOS o PROVEEDOR).")

heading2("7.6 Lo que Ve el Residente vs. lo que Ve el Admin")
_rows_76 = [
    ("Estado general",       "EN ESPERA / EN PROCESO / TERMINADO",           "Lo mismo + días de gestión"),
    ("Ejecución",            "Nota de Fase IV como 'En ejecución'",           "Fases I–V con fechas, días y notas"),
    ("Acción Tomada",        "No visible",                                    "Sí visible y editable"),
    ("Evidencia de cierre",  "Solo texto de evidencia",                       "Texto + archivo adjunto"),
]
_tbl76 = doc.add_table(rows=len(_rows_76)+1, cols=3)
_tbl76.style = "Table Grid"
for _ci, _ht in enumerate(["Elemento", "Vista del Residente", "Vista del Admin/Asistente/Consejo"]):
    _c = _tbl76.cell(0, _ci)
    set_cell_bg(_c, H_VERDE_OSCURO)
    _p = _c.paragraphs[0]
    _r = _p.add_run(_ht)
    _r.font.name = FONT; _r.font.size = Pt(10); _r.font.bold = True; _r.font.color.rgb = BLANCO
    _p.paragraph_format.space_before = Pt(2); _p.paragraph_format.space_after = Pt(2)
for _ri, (_e, _v1, _v2) in enumerate(_rows_76):
    _bg0 = H_VERDE_CLARO if _ri % 2 == 0 else H_BLANCO
    _bg1 = H_GRIS_FONDO if _ri % 2 == 0 else H_BLANCO
    for _ci2, (_txt, _bg) in enumerate([(_e, _bg0), (_v1, _bg1), (_v2, _bg1)]):
        _c2 = _tbl76.cell(_ri+1, _ci2)
        set_cell_bg(_c2, _bg)
        _p2 = _c2.paragraphs[0]
        _r2 = _p2.add_run(_txt)
        _r2.font.name = FONT; _r2.font.size = Pt(9.5)
        if _ci2 == 0: _r2.font.bold = True; _r2.font.color.rgb = VERDE_OSCURO
        _p2.paragraph_format.space_before = Pt(2); _p2.paragraph_format.space_after = Pt(2)
doc.add_paragraph().paragraph_format.space_after = Pt(6)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 8: CIERRE Y EVIDENCIAS
# ═══════════════════════════════════════════════════════════════════════════

heading1("8. Cierre y Evidencias")
body(
    "El cierre de una PQRS documenta formalmente la resolución de la solicitud. "
    "Se puede cerrar una PQRS con o sin completar todas las fases, dependiendo del caso."
)

heading2("8.1 Cierre con Todas las Fases Completadas (Normal)")
body("Cuando se completa la Fase V:")
step_table([
    (1, "Complete la Fase V escribiendo la nota final."),
    (2, "El botón 'Terminar PQRS' aparece disponible en la parte inferior del panel de fases."),
    (3, "Registre la 'Evidencia de Cierre' en el campo de texto (descripción de lo realizado)."),
    (4, "Opcionalmente, adjunte un archivo (imagen, PDF, documento) como evidencia física."),
    (5, "Haga clic en 'Terminar'. El estado cambia a TERMINADO y se registra la fecha de cierre."),
])

heading2("8.2 Cierre Anticipado (Sin Completar Todas las Fases)")
body(
    "En ciertos casos, la PQRS puede resolverse sin seguir todas las fases (por ejemplo, "
    "problemas de convivencia resueltos mediante comunicación)."
)
step_table([
    (1, "En el detalle de la PQRS, ubique el campo '¿Qué se hizo para cerrar?'."),
    (2, "Escriba una descripción detallada de cómo se resolvió la solicitud sin completar las fases."),
    (3, "Cambie el estado a 'TERMINADO' en el selector de estado."),
    (4, "Haga clic en 'Guardar'. El sistema registrará la fecha de cierre automáticamente."),
])

warning_box(
    "El campo '¿Qué se hizo para cerrar?' es obligatorio si se intenta cerrar una PQRS "
    "sin haber completado la Fase V.",
    title="Cierre anticipado"
)

heading2("8.3 Campos de Evidencia de Cierre")
field_table(
    [
        ("Evidencia de cierre",  "Texto descriptivo de lo que se realizó para resolver la solicitud. "
                                  "Puede incluir materiales usados, fechas, proveedores, etc."),
        ("Archivo adjunto",      "Documento, imagen o PDF como evidencia física. "
                                  "Formatos admitidos: JPG, PNG, PDF, DOCX."),
    ],
    header=("Campo", "Descripción")
)

heading2("8.4 Tiempo de Gestión")
body(
    "Al cerrar una PQRS, el sistema calcula automáticamente el tiempo total de gestión: "
    "número de días entre la fecha de recepción y la fecha de cierre. "
    "Este dato es usado en los reportes y estadísticas del Dashboard."
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 9: MÓDULO DE REPORTES
# ═══════════════════════════════════════════════════════════════════════════

heading1("9. Módulo de Reportes")
body(
    "El módulo de reportes permite exportar la información del sistema en formatos Excel (.xlsx) "
    "y PDF para su uso en reuniones, seguimiento y archivo documental."
)

warning_box(
    "El módulo de reportes está disponible exclusivamente para el perfil ADMINISTRADOR.",
    title="Acceso"
)

heading2("9.1 Tipos de Reportes Disponibles")
field_table(
    [
        ("Reporte General PQRS",    "Exporta todas las PQRS del sistema con todos sus campos. "
                                     "Puede filtrarse por año, mes, asunto y estado antes de exportar."),
        ("Reporte por Bloque",       "Resumen de PQRS agrupadas por bloque/torre."),
        ("Reporte por Asunto",       "Resumen de PQRS agrupadas por categoría de asunto."),
        ("Reporte Trimestral",        "Tabla de distribución mensual por trimestre (igual al Dashboard)."),
        ("Reporte de Tiempos",        "Análisis de tiempos de respuesta: primer contacto y cierre."),
    ],
    header=("Tipo", "Contenido")
)

heading2("9.2 Cómo Generar un Reporte Excel")
step_table([
    (1, "Ingrese al módulo 'Reportes' desde el menú lateral."),
    (2, "Seleccione el tipo de reporte que desea generar."),
    (3, "Aplique los filtros deseados (año, mes, asunto, estado)."),
    (4, "Haga clic en el botón 'Exportar Excel'. El archivo .xlsx se descargará automáticamente."),
])

note_box(
    "Los archivos Excel exportados tienen formato profesional con colores, encabezados y "
    "columnas ajustadas para facilitar su presentación.",
    title="Formato Excel"
)

heading2("9.3 Cómo Generar un Reporte PDF")
step_table([
    (1, "Ingrese al módulo 'Reportes' y seleccione los filtros deseados."),
    (2, "Haga clic en el botón 'Exportar PDF'. El navegador descargará el archivo PDF automáticamente."),
])

heading2("9.4 Filtros de Reportes")
field_table(
    [
        ("Año",    "Filtra las PQRS por año de recepción."),
        ("Mes",    "Filtra por mes específico (en combinación con el año)."),
        ("Asunto", "Filtra por categoría de asunto."),
        ("Estado", "Filtra por estado: activas, en espera, en proceso, terminadas o todas."),
    ],
    header=("Filtro", "Descripción")
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 10: GESTIÓN DE USUARIOS
# ═══════════════════════════════════════════════════════════════════════════

heading1("10. Gestión de Usuarios")
body(
    "El módulo de usuarios permite al Administrador crear, editar y gestionar todas las cuentas "
    "del sistema, tanto administrativas como de residentes."
)

warning_box(
    "Solo el perfil ADMINISTRADOR puede acceder al módulo de usuarios y realizar cambios.",
    title="Acceso exclusivo"
)

heading2("10.1 Ver la Lista de Usuarios")
body(
    "En el módulo 'Usuarios' se muestra la lista completa de todos los usuarios registrados. "
    "Se puede filtrar por rol (Admin, Asistente, Consejo, Residente) y buscar por nombre o correo."
)

heading2("10.2 Crear un Usuario Administrativo")
body("Para crear un nuevo perfil de Administrador, Asistente o Consejo:")
step_table([
    (1, "Ingrese al módulo 'Usuarios'."),
    (2, "Haga clic en 'Nuevo Usuario'."),
    (3, "Complete el nombre completo, correo electrónico y contraseña inicial."),
    (4, "Seleccione el rol: ADMIN, ASISTENTE o CONSEJO."),
    (5, "Haga clic en 'Crear'. El usuario podrá ingresar de inmediato con esas credenciales."),
])

warning_box(
    "Entregue la contraseña inicial al nuevo usuario de forma segura y solicítele que la cambie "
    "al primer ingreso desde la opción 'Cambiar contraseña'.",
    title="Seguridad"
)

heading2("10.3 Crear un Usuario Residente")
body("Para crear una cuenta de residente:")
step_table([
    (1, "En el módulo 'Usuarios', haga clic en 'Nuevo Usuario'."),
    (2, "Complete el nombre completo, correo electrónico y contraseña."),
    (3, "Seleccione el rol 'RESIDENTE'."),
    (4, "Ingrese el Bloque (1–12) y el número de Apartamento del residente."),
    (5, "Haga clic en 'Crear'."),
])

info_box(
    "Los residentes también pueden registrarse de forma autónoma mediante el código QR del conjunto. "
    "El QR redirige a la página de registro donde el residente completa sus datos.",
    title="Registro por QR"
)

heading2("10.4 Editar un Usuario")
step_table([
    (1, "En la lista de usuarios, haga clic sobre el usuario que desea editar."),
    (2, "Modifique los campos necesarios: nombre, bloque, apartamento o rol."),
    (3, "Para cambiar la contraseña de otro usuario, use el campo 'Nueva contraseña'."),
    (4, "Haga clic en 'Guardar cambios'."),
])

heading2("10.5 Desactivar o Eliminar un Usuario")
body(
    "Si un residente ya no vive en el conjunto o un colaborador deja sus funciones, "
    "el Administrador puede eliminar su cuenta desde el módulo de usuarios. "
    "Al eliminar un usuario, sus PQRS históricas permanecen registradas en el sistema."
)

heading2("10.6 Campos del Perfil de Usuario")
field_table(
    [
        ("Nombre",      "Nombre completo del usuario."),
        ("Correo",      "Correo electrónico (usado para iniciar sesión). Debe ser único."),
        ("Contraseña",  "Contraseña de acceso (encriptada en el sistema)."),
        ("Rol",         "ADMIN, ASISTENTE, CONSEJO o RESIDENTE."),
        ("Bloque",      "Número de bloque (solo para residentes). Del 1 al 12."),
        ("Apartamento", "Número de apartamento (solo para residentes)."),
    ],
    header=("Campo", "Descripción")
)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# SECCIÓN 11: PREGUNTAS FRECUENTES
# ═══════════════════════════════════════════════════════════════════════════

heading1("11. Preguntas Frecuentes")

faqs = [
    (
        "¿Puede un Asistente cerrar una PQRS?",
        "Sí. El Asistente puede cambiar el estado a TERMINADO y registrar evidencia de cierre. "
        "Solo la creación de usuarios y el acceso a reportes están limitados al Administrador."
    ),
    (
        "¿El residente puede ver las fases de gestión?",
        "No directamente. El residente solo ve el estado general (EN ESPERA, EN PROCESO, TERMINADO) "
        "y, cuando está en proceso, puede ver la nota de la Fase IV bajo el título 'En ejecución'."
    ),
    (
        "¿Qué pasa si me equivoqué al avanzar una fase?",
        "Contacte al Administrador del sistema. No existe función de retroceder fases automáticamente, "
        "pero el Administrador puede editar directamente los campos de fase si es necesario."
    ),
    (
        "¿Puedo cerrar una PQRS sin completar todas las fases?",
        "Sí. Utilice el campo '¿Qué se hizo para cerrar?' para documentar la resolución. "
        "Este campo es obligatorio cuando se cierra sin completar la Fase V."
    ),
    (
        "¿Cómo sé cuánto tiempo lleva abierta una PQRS?",
        "En el Dashboard y en la lista de PQRS activas, cada solicitud muestra el número de días "
        "transcurridos. En el detalle, se muestra la fecha exacta de recepción."
    ),
    (
        "¿Qué formatos acepta el archivo de evidencia?",
        "El sistema acepta imágenes (JPG, PNG), documentos PDF y archivos Word (DOCX). "
        "El tamaño máximo recomendado es 5 MB."
    ),
    (
        "¿Puedo editar una PQRS después de crearla?",
        "Sí. Los campos de gestión (tipo, asunto, primer contacto, acción tomada, estado) "
        "pueden editarse en cualquier momento desde el detalle de la PQRS."
    ),
    (
        "¿Cómo funciona el registro por QR para residentes?",
        "El Administrador puede compartir el código QR del conjunto. Al escanearlo, el residente "
        "llega a la página de registro donde completa sus datos personales y crea su cuenta. "
        "El rol RESIDENTE se asigna automáticamente."
    ),
    (
        "¿Los reportes incluyen PQRS de todos los años?",
        "Sí. Por defecto los reportes incluyen todas las PQRS. Use el filtro de año "
        "para restringir el período del reporte."
    ),
    (
        "¿Qué significa el indicador de semáforo en las fases?",
        "Verde (0-7 días): gestión en tiempo normal. Amarillo (8-14 días): requiere atención pronto. "
        "Rojo (15+ días): la fase está demorada y requiere acción inmediata."
    ),
]

for i, (q, a) in enumerate(faqs):
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"
    cq = tbl.cell(0, 0); ca = tbl.cell(1, 0)
    set_cell_bg(cq, H_VERDE_MEDIO); set_cell_bg(ca, H_GRIS_FONDO)
    pq = cq.paragraphs[0]
    rq = pq.add_run(f"P{i+1}. {q}")
    rq.font.name = FONT; rq.font.size = Pt(10); rq.font.bold = True; rq.font.color.rgb = BLANCO
    pq.paragraph_format.space_before = Pt(3); pq.paragraph_format.space_after = Pt(3)
    pa = ca.paragraphs[0]
    ra = pa.add_run(a)
    ra.font.name = FONT; ra.font.size = Pt(10)
    pa.paragraph_format.space_before = Pt(3); pa.paragraph_format.space_after = Pt(3)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

add_page_break()

# ═══════════════════════════════════════════════════════════════════════════
# PÁGINA FINAL
# ═══════════════════════════════════════════════════════════════════════════

doc.add_paragraph()
doc.add_paragraph()
doc.add_paragraph()

tbl_end = doc.add_table(rows=1, cols=1)
tbl_end.style = "Table Grid"
c_end = tbl_end.cell(0, 0)
set_cell_bg(c_end, H_VERDE_OSCURO)
p_end = c_end.paragraphs[0]
r_end = p_end.add_run("CONJUNTO RESIDENCIAL CALLE 100\nSistema PQRS — Versión 1.0 — 2026\nDocumento de uso interno")
r_end.font.name = FONT; r_end.font.size = Pt(11); r_end.font.bold = True; r_end.font.color.rgb = BLANCO
p_end.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_end.paragraph_format.space_before = Pt(16); p_end.paragraph_format.space_after = Pt(16)

# ─── Guardar ───────────────────────────────────────────────────────────────
out_path = "Manual_Admin_PQRS.docx"
doc.save(out_path)
print(f"Manual Admin generado: {out_path}")
