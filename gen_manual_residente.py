from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

GREEN = RGBColor(0x1B, 0x5E, 0x20)
DARK  = RGBColor(0x21, 0x21, 0x21)
GRAY  = RGBColor(0x55, 0x55, 0x55)
LIGHT_GREEN = RGBColor(0x2E, 0x7D, 0x32)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def heading1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = GREEN
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    '8')
    bot.set(qn('w:space'), '2')
    bot.set(qn('w:color'), '1B5E20')
    pBdr.append(bot)
    pPr.append(pBdr)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = LIGHT_GREEN
    return p

def body(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(5)
    for run in p.runs:
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p

def bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_part and bold_part in text:
        idx = text.index(bold_part)
        before = text[:idx]
        after  = text[idx+len(bold_part):]
        if before:
            r = p.add_run(before); r.font.size = Pt(11); r.font.color.rgb = DARK
        r2 = p.add_run(bold_part); r2.bold = True; r2.font.size = Pt(11); r2.font.color.rgb = DARK
        if after:
            r3 = p.add_run(after); r3.font.size = Pt(11); r3.font.color.rgb = DARK
    else:
        run = p.add_run(text)
        run.font.size = Pt(11)
        run.font.color.rgb = DARK
    return p

def note_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'E8F5E9')
    p = cell.paragraphs[0]
    run = p.add_run('💡  ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x1B, 0x5E, 0x20)
    run.italic = True
    doc.add_paragraph()

def warning_box(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.cell(0, 0)
    set_cell_bg(cell, 'FFF3E0')
    p = cell.paragraphs[0]
    run = p.add_run('⚠️  ' + text)
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0xE6, 0x5C, 0x00)
    run.italic = True
    doc.add_paragraph()

def step_table(doc, steps):
    """steps: list of (numero, titulo, descripcion)"""
    tbl = doc.add_table(rows=len(steps), cols=2)
    tbl.style = 'Table Grid'
    tbl.columns[0].width = Cm(1.5)
    tbl.columns[1].width = Cm(14)
    for i, (num, title, desc) in enumerate(steps):
        c0 = tbl.cell(i, 0)
        c1 = tbl.cell(i, 1)
        set_cell_bg(c0, '2E7D32')
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(num))
        r0.bold = True; r0.font.size = Pt(14); r0.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        c1.paragraphs[0].clear()
        p1 = c1.paragraphs[0]
        r_title = p1.add_run(title + '\n')
        r_title.bold = True; r_title.font.size = Pt(11); r_title.font.color.rgb = DARK
        r_desc = p1.add_run(desc)
        r_desc.font.size = Pt(10); r_desc.font.color.rgb = GRAY
    doc.add_paragraph()

def status_table(doc):
    tbl = doc.add_table(rows=4, cols=3)
    tbl.style = 'Table Grid'
    # header
    hdrs = ['Estado', 'Significado', 'Qué puede hacer el residente']
    fills = ['2E7D32','2E7D32','2E7D32']
    for i, h in enumerate(hdrs):
        c = tbl.cell(0, i)
        set_cell_bg(c, '1B5E20')
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(255,255,255)

    rows_data = [
        ('⏳  En Espera',    'FFF9C4', 'La solicitud fue recibida y está pendiente de revisión por parte de la administración.',
         'Consultar el detalle. No se requiere acción adicional.'),
        ('🕐  En Proceso',   'E3F2FD', 'La administración está gestionando activamente la solicitud.',
         'Consultar el avance. Si hay ejecución, se mostrará la sección "En ejecución".'),
        ('✅  Terminado',    'E8F5E9', 'La solicitud fue resuelta y cerrada por la administración.',
         'Consultar el resumen del cierre y la evidencia adjunta.'),
    ]
    for i, (estado, fill, sig, accion) in enumerate(rows_data, 1):
        tbl.cell(i, 0).paragraphs[0].add_run(estado).font.size = Pt(10)
        set_cell_bg(tbl.cell(i, 0), fill)
        tbl.cell(i, 1).paragraphs[0].add_run(sig).font.size = Pt(10)
        tbl.cell(i, 2).paragraphs[0].add_run(accion).font.size = Pt(10)
    doc.add_paragraph()

# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# Page setup
for sec in doc.sections:
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin  = sec.bottom_margin = Cm(2.5)

# ── PORTADA ──────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run('PORTAL WEB PQRS\nCONJUNTO PARQUE RESIDENCIAL CALLE 100')
r.bold = True; r.font.size = Pt(24); r.font.color.rgb = GREEN

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('Manual de Usuario — Residente')
r2.font.size = Pt(16); r2.font.color.rgb = GRAY

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Versión 1.0  ·  Abril 2026')
r3.font.size = Pt(11); r3.font.color.rgb = GRAY

doc.add_page_break()

# ── ÍNDICE ───────────────────────────────────────────────────────────────────
heading1(doc, 'Contenido')
toc_items = [
    ('1.', 'Introducción'),
    ('2.', 'Acceso al portal — Inicio de sesión'),
    ('3.', 'Pantalla principal — Mis PQRS'),
    ('4.', 'Crear una nueva PQRS'),
    ('5.', 'Detalle de una PQRS'),
    ('6.', 'Estados de una PQRS'),
    ('7.', 'Seguimiento durante el proceso'),
    ('8.', 'Cambiar contraseña'),
    ('9.', 'Cerrar sesión'),
    ('10.','Preguntas frecuentes'),
]
for num, title in toc_items:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r1 = p.add_run(f'{num}  ')
    r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = GREEN
    r2 = p.add_run(title)
    r2.font.size = Pt(11); r2.font.color.rgb = DARK

doc.add_page_break()

# ── 1. INTRODUCCIÓN ───────────────────────────────────────────────────────────
heading1(doc, '1.  Introducción')
body(doc,
    'El Portal Web PQRS del Conjunto Parque Residencial Calle 100 es la plataforma '
    'digital oficial para que los residentes puedan radicar, consultar y hacer seguimiento '
    'a sus Peticiones, Quejas, Reclamos y Sugerencias (PQRS) directamente desde su '
    'computador o dispositivo móvil, sin necesidad de acercarse a la administración.')
body(doc,
    'Este manual explica paso a paso cómo utilizar todas las funcionalidades '
    'disponibles para el rol de Residente.')
note_box(doc,
    'Para acceder al portal necesita las credenciales (correo electrónico y contraseña) '
    'que le fueron asignadas por la administración del conjunto.')

# ── 2. ACCESO AL PORTAL ───────────────────────────────────────────────────────
heading1(doc, '2.  Acceso al portal — Inicio de sesión')
body(doc, 'Siga los siguientes pasos para ingresar al portal:')
step_table(doc, [
    (1, 'Abrir el navegador web',
     'Ingrese a la dirección del portal proporcionada por la administración '
     '(Google Chrome, Firefox o Edge recomendados).'),
    (2, 'Ingresar credenciales',
     'En la pantalla de inicio de sesión escriba su correo electrónico y su contraseña.'),
    (3, 'Hacer clic en "Iniciar sesión"',
     'Si los datos son correctos, el sistema lo redirigirá automáticamente a la pantalla '
     'principal con su listado de PQRS.'),
])
warning_box(doc,
    'Si olvidó su contraseña, contacte a la administración para que le genere unas '
    'nuevas credenciales de acceso.')

# ── 3. PANTALLA PRINCIPAL ─────────────────────────────────────────────────────
heading1(doc, '3.  Pantalla principal — Mis PQRS')
body(doc,
    'Una vez dentro del portal, verá la sección "Mis PQRS", que muestra el listado '
    'completo de todas las solicitudes que usted ha radicado.')

heading2(doc, '3.1  Elementos de la pantalla')
bullet(doc, 'Título "Mis PQRS": confirma que está viendo únicamente sus solicitudes.', 'Título "Mis PQRS"')
bullet(doc, 'Botón "Crear": permite radicar una nueva PQRS (esquina superior derecha).', 'Botón "Crear"')
bullet(doc, 'Contador de resultados: indica cuántas solicitudes coinciden con los filtros activos.', 'Contador de resultados')

heading2(doc, '3.2  Filtros disponibles')
body(doc, 'Puede combinar los siguientes filtros para encontrar una solicitud específica:')

tbl = doc.add_table(rows=4, cols=2)
tbl.style = 'Table Grid'
for i, (f, d) in enumerate([
    ('Filtro', 'Descripción'),
    ('Estado',  'Filtra por En Espera, En Proceso o Terminadas.'),
    ('Mes',     'Filtra solicitudes por el mes en que fueron radicadas.'),
    ('Año',     'Filtra solicitudes por el año de radicación.'),
]):
    set_cell_bg(tbl.cell(i,0), '1B5E20' if i==0 else 'F1F8E9')
    tbl.cell(i,0).paragraphs[0].add_run(f).bold = (i==0)
    tbl.cell(i,0).paragraphs[0].runs[-1].font.size = Pt(11)
    tbl.cell(i,0).paragraphs[0].runs[-1].font.color.rgb = RGBColor(255,255,255) if i==0 else DARK
    tbl.cell(i,1).paragraphs[0].add_run(d).font.size = Pt(11)
doc.add_paragraph()

note_box(doc,
    'Use el botón "Limpiar" para restablecer todos los filtros y ver el listado completo.')

heading2(doc, '3.3  Tarjeta de resumen de una PQRS')
body(doc,
    'Cada solicitud aparece como una tarjeta con la siguiente información resumida:')
bullet(doc, 'Número de radicado (#001, #002, etc.)')
bullet(doc, 'Chip de estado (color amarillo = En Espera, azul = En Proceso, verde = Terminado)')
bullet(doc, 'Asunto o primeras palabras de la descripción')
bullet(doc, 'Fecha de radicación')
body(doc, 'Haga clic sobre cualquier tarjeta para ver el detalle completo.')

# ── 4. CREAR UNA PQRS ─────────────────────────────────────────────────────────
heading1(doc, '4.  Crear una nueva PQRS')
body(doc,
    'Para radicar una nueva solicitud, haga clic en el botón verde "Crear" ubicado '
    'en la esquina superior derecha de la pantalla "Mis PQRS".')

step_table(doc, [
    (1, 'Seleccionar el Asunto (opcional)',
     'Escoja la categoría que mejor describe su solicitud: Área Común, Contabilidad, '
     'Convivencia, Humedad/Cubierta, Humedad/Depósito, Humedad/Ventanas, '
     'Humedad/Fachada o Humedad/Garaje.'),
    (2, 'Redactar la descripción',
     'Escriba con claridad el detalle de su petición, queja, reclamo o sugerencia. '
     'El límite es de 300 palabras. Un contador en tiempo real muestra las palabras '
     'utilizadas (se torna rojo al superar el límite).'),
    (3, 'Enviar la solicitud',
     'Haga clic en el botón "Enviar solicitud". El sistema asignará automáticamente '
     'un número de radicado y la PQRS quedará en estado "En Espera".'),
])
note_box(doc,
    'Sus datos personales (nombre, bloque y apartamento) se asocian automáticamente '
    'a la solicitud a partir de su cuenta. No es necesario ingresarlos manualmente.')
warning_box(doc,
    'La descripción es el único campo obligatorio. Si supera 300 palabras, el sistema '
    'impedirá el envío hasta que reduzca el texto.')

# ── 5. DETALLE DE UNA PQRS ────────────────────────────────────────────────────
heading1(doc, '5.  Detalle de una PQRS')
body(doc,
    'Al hacer clic sobre una tarjeta de la lista, se abre la vista de detalle con toda '
    'la información relacionada a esa solicitud.')

heading2(doc, '5.1  Información general')
bullet(doc, 'Número de radicado y estado actual')
bullet(doc, 'Fecha de recibido y mes')
bullet(doc, 'Bloque y apartamento')
bullet(doc, 'Nombre del residente')
bullet(doc, 'Asunto y descripción completa de la solicitud')

heading2(doc, '5.2  Información de gestión (visible según el estado)')
body(doc,
    'A medida que la administración avanza en la gestión, aparecerán nuevas secciones:')
bullet(doc, 'Fecha de primer contacto y tiempo de respuesta inicial', 'Fecha de primer contacto')
bullet(doc, 'Sección "En ejecución": disponible cuando la solicitud está siendo ejecutada '
           '(ver Sección 7).', '"En ejecución"')
bullet(doc, 'Evidencia de cierre: texto o archivo adjunto que describe la solución aplicada.', 'Evidencia de cierre')
bullet(doc, 'Fecha de cierre y tiempo total de respuesta.', 'Fecha de cierre')

# ── 6. ESTADOS ────────────────────────────────────────────────────────────────
heading1(doc, '6.  Estados de una PQRS')
body(doc,
    'Toda PQRS pasa por tres estados a lo largo de su ciclo de vida. '
    'El estado se muestra como un chip de color tanto en la lista como en el detalle.')
status_table(doc)
note_box(doc,
    'El residente no puede cambiar manualmente el estado. Los cambios los realiza '
    'exclusivamente la administración del conjunto.')

# ── 7. SEGUIMIENTO EN PROCESO ─────────────────────────────────────────────────
heading1(doc, '7.  Seguimiento durante el proceso')
body(doc,
    'Cuando su PQRS se encuentre en estado "En Proceso", la administración puede '
    'registrar notas de avance por cada fase de gestión interna.')
heading2(doc, '7.1  Sección "En ejecución"')
body(doc,
    'Si la administración ha iniciado la fase de ejecución de obras o acciones y ha '
    'registrado una nota al respecto, usted verá automáticamente en el detalle de su '
    'PQRS una tarjeta azul titulada "En ejecución" con la descripción de lo que se '
    'está haciendo.')
note_box(doc,
    'Esta información es de solo lectura. Refleja en tiempo real el progreso registrado '
    'por la administración sin revelar la estructura interna de fases de gestión.')
heading2(doc, '7.2  ¿Qué no puede ver el residente?')
body(doc,
    'Para proteger la confidencialidad operativa del proceso, las siguientes '
    'secciones no son visibles para el residente:')
bullet(doc, 'Detalle de las fases de gestión interna (Fase I a V)')
bullet(doc, 'Acción tomada por el gestor')
bullet(doc, 'Notas internas de cada fase')

# ── 8. CAMBIAR CONTRASEÑA ─────────────────────────────────────────────────────
heading1(doc, '8.  Cambiar contraseña')
body(doc,
    'Si desea actualizar su contraseña de acceso, contacte a la administración del '
    'conjunto para que le asigne unas nuevas credenciales. Actualmente el cambio de '
    'contraseña es gestionado por el administrador del sistema.')
note_box(doc,
    'Mantenga su contraseña confidencial y no la comparta con terceros. Si sospecha '
    'que su cuenta fue comprometida, notifique de inmediato a la administración.')

# ── 9. CERRAR SESIÓN ──────────────────────────────────────────────────────────
heading1(doc, '9.  Cerrar sesión')
step_table(doc, [
    (1, 'Localizar el menú de usuario',
     'En la esquina superior derecha de cualquier pantalla encontrará su nombre o '
     'un ícono de perfil.'),
    (2, 'Seleccionar "Cerrar sesión"',
     'Haga clic en la opción correspondiente. El sistema lo redirigirá a la pantalla '
     'de inicio de sesión.'),
])
warning_box(doc,
    'Siempre cierre sesión cuando utilice un dispositivo compartido o público.')

# ── 10. PREGUNTAS FRECUENTES ──────────────────────────────────────────────────
heading1(doc, '10.  Preguntas frecuentes')

faqs = [
    ('¿Puedo ver las PQRS de otros residentes?',
     'No. El sistema filtra automáticamente y le muestra únicamente las solicitudes '
     'asociadas a su cuenta.'),
    ('¿Puedo modificar o eliminar una PQRS ya enviada?',
     'No. Una vez radicada, la PQRS no puede ser editada ni eliminada por el residente. '
     'Si necesita una corrección, contacte a la administración.'),
    ('¿Recibiré una notificación cuando mi PQRS cambie de estado?',
     'El sistema puede enviar notificaciones por correo electrónico cuando la administración '
     'actualice el estado de su solicitud, según la configuración del conjunto.'),
    ('¿Cuántas PQRS puedo radicar?',
     'No existe un límite definido de solicitudes por residente.'),
    ('¿La descripción tiene límite de extensión?',
     'Sí. El campo de descripción admite un máximo de 300 palabras. Un contador en '
     'tiempo real le indica cuántas ha utilizado.'),
    ('¿Qué hago si no puedo iniciar sesión?',
     'Verifique que su correo y contraseña sean correctos. Si el problema persiste, '
     'contacte a la administración para restablecer sus credenciales.'),
]
for q, a in faqs:
    heading2(doc, q)
    body(doc, a)

# ── PIE ───────────────────────────────────────────────────────────────────────
doc.add_page_break()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Conjunto Parque Residencial Calle 100  ·  Portal PQRS  ·  Versión 1.0  ·  2026')
r.font.size = Pt(9); r.font.color.rgb = GRAY; r.italic = True

doc.save(r'c:\Users\d.hernandeza2\Documents\PQRS_CALLE_100\Manual_Residente_PQRS.docx')
print("Manual Residente generado.")
